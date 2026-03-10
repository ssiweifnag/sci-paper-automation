from __future__ import annotations

import re
from typing import Any, Dict


class FormatChecker:
    REQUIRED_SECTIONS = [
        'abstract', 'introduction', 'methods', 'results', 'discussion', 'conclusion', 'references'
    ]

    def check_docx(self, filepath: str) -> Dict[str, Any]:
        try:
            from docx import Document  # type: ignore
        except ImportError:
            return {'skipped': True, 'reason': '未安裝 python-docx，暫時跳過 DOCX 格式檢查'}

        doc = Document(filepath)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        headings = [
            p.text.strip() for p in doc.paragraphs
            if getattr(p.style, 'name', '').lower().startswith('heading') and p.text.strip()
        ]
        text = '\n'.join(paragraphs)
        return {
            'structure': self._check_structure(headings),
            'citations': self._check_citations(text),
            'figures': self._check_figures(paragraphs),
            'word_count': len(text.split()),
        }

    def _check_structure(self, headings: list[str]) -> Dict[str, Any]:
        found = [h.lower() for h in headings]
        missing = [s for s in self.REQUIRED_SECTIONS if not any(s in h for h in found)]
        return {'missing_sections': missing, 'found_sections': headings}

    def _check_citations(self, text: str) -> Dict[str, Any]:
        apa_pattern = r'\([A-Z][A-Za-z\-]+(?:\s+et\s+al\.)?,\s*\d{4}\)'
        vancouver_pattern = r'\[\d+\]'
        apa_count = len(re.findall(apa_pattern, text))
        van_count = len(re.findall(vancouver_pattern, text))
        return {
            'detected_style': 'APA' if apa_count >= van_count else 'Vancouver',
            'apa_citations': apa_count,
            'numbered_citations': van_count,
            'consistent': not (apa_count > 0 and van_count > 0),
        }

    def _check_figures(self, paragraphs: list[str]) -> Dict[str, Any]:
        captions = [p for p in paragraphs if p.startswith('Figure') or p.startswith('Fig.')]
        return {'figure_captions': captions, 'count': len(captions)}
